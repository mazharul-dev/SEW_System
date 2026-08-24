from __future__ import annotations

from io import BytesIO
import re
from zipfile import ZipFile

import pytest
from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

from app import app


UNICODE_BANGLA_RE = re.compile(r"[\u0964\u0965\u0980-\u09ff]")


@pytest.fixture(autouse=True)
def _isolated_join_store(tmp_path):
    app.config["JOIN_REQUESTS_FILE"] = str(tmp_path / "join_requests.json")
    yield


def _join_and_approve(client, email="user@example.com"):
    response = client.post("/join/request", json={"email": email})
    assert response.status_code == 200
    response = client.post("/admin/login", data={"username": "admin", "password": "01706452007"})
    assert response.status_code == 302
    response = client.post("/admin/approvals/approve", data={"email": email})
    assert response.status_code == 302


def _table_mcq_docx() -> BytesIO:
    return _mcq_docx("1", "General", "What is 2 + 2?")


def _mcq_docx(serial: str, category: str, question: str) -> BytesIO:
    document = Document()
    table = document.add_table(rows=8, cols=2)
    table.cell(0, 0).text = serial
    table.cell(0, 1).text = category
    table.cell(1, 0).text = question
    table.cell(2, 0).text = "3"
    table.cell(3, 0).text = "4"
    table.cell(4, 0).text = "5"
    table.cell(5, 0).text = "6"
    table.cell(6, 0).text = "Two plus two equals four."
    table.cell(7, 0).text = "B"
    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer


def _plain_mcq_docx() -> BytesIO:
    document = Document()
    document.add_paragraph("1. What is 2 + 2?")
    document.add_paragraph("(a) 3")
    document.add_paragraph("(b) 4")
    document.add_paragraph("(c) 5")
    document.add_paragraph("(d) 6")
    document.add_paragraph("Answer: b")
    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer


def _unicode_question_docx() -> BytesIO:
    document = Document()
    document.add_paragraph("\u09ac\u09be\u0982\u09b2\u09be \u09aa\u09cd\u09b0\u09b6\u09cd\u09a8\u0964")
    document.sections[0].header.paragraphs[0].text = "\u09b9\u09c7\u09a1\u09be\u09b0 \u09b2\u09c7\u0996\u09be"
    document.sections[0].footer.paragraphs[0].text = "\u09ab\u09c1\u099f\u09be\u09b0 \u09b2\u09c7\u0996\u09be"
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "\u099c\u09c0\u09ac\u09a8\u09be\u09a8\u09a8\u09cd\u09a6 \u09a6\u09be\u09b6"
    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer


def _mixed_math_image_docx() -> BytesIO:
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("\u09ac\u09be\u0982\u09b2\u09be ")
    paragraph._p.append(parse_xml(f'<m:oMath {nsdecls("m")}><m:r><m:t>x</m:t></m:r></m:oMath>'))
    paragraph.add_run(" \u09aa\u09cd\u09b0\u09b6\u09cd\u09a8")

    mixed_run_paragraph = document.add_paragraph()
    mixed_run_paragraph._p.append(
        parse_xml(
            f'<w:r {nsdecls("w", "m")}>'
            "<w:t>\u09b9\u09b2\u09c7 </w:t>"
            "<m:oMath><m:r><m:t>m = n</m:t></m:r></m:oMath>"
            "<w:t> \u098f\u09ac\u0982 \u09aa\u09cd\u09b0\u09b6\u09cd\u09a8</w:t>"
            "</w:r>"
        )
    )

    cell_paragraph = document.add_table(rows=1, cols=1).cell(0, 0).paragraphs[0]
    cell_paragraph.add_run("\u099b\u09ac\u09bf ")
    cell_paragraph._p.append(parse_xml(f'<w:r {nsdecls("w")}><w:drawing /></w:r>'))
    cell_paragraph.add_run(" \u09b2\u09c7\u0996\u09be")

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer


def _plain_math_inside_bangla_docx() -> BytesIO:
    document = Document()
    document.add_paragraph(
        "\u09b9\u09b2\u09c7, m = n \u098f\u09ac\u0982 |x-2| < 3 "
        "\u098f\u09b0 \u09ae\u09be\u09a8 \u0995\u09a4? English\u0964"
    )
    document.add_paragraph("Maths English 123")
    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer


def test_dashboard_and_tool_routes_load():
    client = app.test_client()
    _join_and_approve(client)
    for path in (
        "/",
        "/tools/bangla-convert",
        "/tools/model-test-generate",
        "/tools/question-proofreader",
        "/tools/table-converter",
        "/tools/in-branch-question-convert",
        "/tools/multi-question-set-repeat-checker",
        "/tools/bangla-convert/_module/",
        "/tools/model-test-generate/_module/",
        "/tools/question-proofreader/_module/",
        "/admin/approvals",
    ):
        assert client.get(path).status_code == 200


def test_proofreader_parse_api_accepts_table_docx():
    client = app.test_client()
    _join_and_approve(client)
    response = client.post(
        "/tools/question-proofreader/api/parse",
        data={"upload": (_table_mcq_docx(), "questions.docx")},
        content_type="multipart/form-data",
    )

    assert response.status_code == 200
    payload = response.get_json()
    assert payload["total"] == 1
    assert payload["questions"][0]["answerLabel"]


def test_table_converter_downloads_docx():
    client = app.test_client()
    _join_and_approve(client)
    response = client.post(
        "/tools/table-converter/convert",
        data={"subject": "General", "upload": (_plain_mcq_docx(), "plain.docx")},
        content_type="multipart/form-data",
    )

    assert response.status_code == 200
    assert response.mimetype == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def test_model_test_generator_downloads_docx():
    client = app.test_client()
    _join_and_approve(client)
    response = client.post(
        "/tools/model-test-generate/convert",
        data={
            "delete_solve": "on",
            "convert_omr": "on",
            "two_column": "on",
            "upload": (_plain_mcq_docx(), "model.docx"),
        },
        content_type="multipart/form-data",
    )

    assert response.status_code == 200
    assert response.mimetype == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def test_in_branch_converter_downloads_sutonny_docx():
    client = app.test_client()
    _join_and_approve(client)
    response = client.post(
        "/tools/in-branch-question-convert/convert",
        data={"upload": (_unicode_question_docx(), "branch.docx")},
        content_type="multipart/form-data",
    )

    assert response.status_code == 200
    assert response.mimetype == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    output = Document(BytesIO(response.data))
    body_text = output.paragraphs[0].text
    table_text = output.tables[0].cell(0, 0).text
    header_text = output.sections[0].header.paragraphs[0].text
    footer_text = output.sections[0].footer.paragraphs[0].text
    assert "|" in body_text
    assert not UNICODE_BANGLA_RE.search(body_text + table_text + header_text + footer_text)
    assert output.paragraphs[0].runs[0].font.name == "SutonnyMJ"
    assert output.tables[0].cell(0, 0).paragraphs[0].runs[0].font.name == "SutonnyMJ"
    assert output.sections[0].header.paragraphs[0].runs[0].font.name == "SutonnyMJ"
    assert output.sections[0].footer.paragraphs[0].runs[0].font.name == "SutonnyMJ"


def test_in_branch_converter_preserves_math_and_drawing_xml():
    client = app.test_client()
    _join_and_approve(client)
    source = _mixed_math_image_docx().getvalue()
    response = client.post(
        "/tools/in-branch-question-convert/convert",
        data={"upload": (BytesIO(source), "mixed.docx")},
        content_type="multipart/form-data",
    )

    assert response.status_code == 200

    input_xml = ZipFile(BytesIO(source)).read("word/document.xml").decode("utf-8")
    output_xml = ZipFile(BytesIO(response.data)).read("word/document.xml").decode("utf-8")
    assert input_xml.count("<m:oMath") == output_xml.count("<m:oMath")
    assert input_xml.count("<w:drawing") == output_xml.count("<w:drawing")
    assert not UNICODE_BANGLA_RE.search(output_xml)


def test_in_branch_converter_preserves_plain_math_inside_bangla_text():
    client = app.test_client()
    _join_and_approve(client)
    response = client.post(
        "/tools/in-branch-question-convert/convert",
        data={"upload": (_plain_math_inside_bangla_docx(), "plain-math.docx")},
        content_type="multipart/form-data",
    )

    assert response.status_code == 200

    output = Document(BytesIO(response.data))
    paragraph = output.paragraphs[0]
    preserved_runs = [run for run in paragraph.runs if "m = n" in run.text or "|x-2| < 3" in run.text]
    assert len(preserved_runs) == 2
    assert all(run.font.name == "Times New Roman" for run in preserved_runs)
    assert "m = n" in paragraph.text
    assert "|x-2| < 3" in paragraph.text
    assert "English" in paragraph.text
    assert "English|" in paragraph.text
    assert not UNICODE_BANGLA_RE.search(paragraph.text)

    english_paragraph = output.paragraphs[1]
    assert english_paragraph.runs[0].font.name == "Times New Roman"


def test_multi_question_set_repeat_checker_reports_cross_set_repeats():
    client = app.test_client()
    _join_and_approve(client)
    response = client.post(
        "/tools/multi-question-set-repeat-checker",
        data={
            "daily_live": (_mcq_docx("1", "Bangla", "Which poet wrote Banalata Sen?"), "daily-live.docx"),
            "daily_practice": (
                _mcq_docx("12", "Bangla", "Which poet wrote Banalata Sen?"),
                "daily-practice.docx",
            ),
            "weekly_live": (_mcq_docx("22", "Bangla", "Who wrote Gitanjali?"), "weekly-live.docx"),
            "weekly_practice": (
                _mcq_docx("33", "Bangla", "Which poet wrote Banalata Sen?"),
                "weekly-practice.docx",
            ),
        },
        content_type="multipart/form-data",
    )

    assert response.status_code == 200
    html = response.get_data(as_text=True)
    assert "Cross Set Repeat" in html
    assert "Daily Live" in html
    assert "Daily Practice" in html
    assert "Weekly Practice" in html
    assert "Serial 12" in html


def test_tool_routes_are_blocked_before_join_approval():
    client = app.test_client()
    response = client.get("/tools/bangla-convert")
    assert response.status_code == 302


def test_admin_approval_panel_requires_login():
    client = app.test_client()
    response = client.get("/admin/approvals")
    assert response.status_code == 302
    assert "/admin/login" in response.headers["Location"]


def test_previously_approved_email_unlocks_without_second_admin_approval():
    first_client = app.test_client()
    _join_and_approve(first_client, email="repeat@example.com")

    second_client = app.test_client()
    response = second_client.post("/join/request", json={"email": "repeat@example.com"})
    assert response.status_code == 200
    payload = response.get_json()
    assert payload["access"]["approved"] is True
    assert payload["request"]["status"] == "approved"

    response = second_client.get("/tools/bangla-convert")
    assert response.status_code == 200
