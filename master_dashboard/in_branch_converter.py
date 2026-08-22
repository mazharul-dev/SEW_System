from __future__ import annotations

from collections.abc import Iterable
from copy import deepcopy
from dataclasses import dataclass
from io import BytesIO
import json
import re
import subprocess
from typing import Any

from docx import Document
from docx.document import Document as DocumentObject
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.text.run import Run

from .paths import MAZHARUL_CONVERTER_DIR, PACKAGE_DIR


SUTONNY_FONT = "SutonnyMJ"
ENGLISH_FONT = "Times New Roman"
BANGLA_RE = re.compile(r"[\u0964\u0965\u0980-\u09ff]")
CONVERTER_SCRIPT = PACKAGE_DIR / "converters" / "unicode_to_bijoy_batch.js"
FONTCONVERTER_SCRIPT = MAZHARUL_CONVERTER_DIR / "js" / "fc" / "fontconverter.min.js"
XML_SPACE = "{http://www.w3.org/XML/1998/namespace}space"
TEXT_LOCAL_NAMES = {"t", "tab", "br", "cr"}
EMBEDDED_LOCAL_NAMES = {"drawing", "pict", "object", "oMath", "oMathPara"}
NON_BANGLA_CHUNK_RE = re.compile(r"[^\u0964\u0965\u0980-\u09ff]+")
MATH_OPERATOR_CHARS = set("=+-*/^<>") | {
    "\u00d7",
    "\u00f7",
    "\u2212",
    "\u221a",
    "\u221e",
    "\u2260",
    "\u2264",
    "\u2265",
    "\u2248",
}
WORD_TEXT_TAG = qn("w:t")
WORD_TAB_TAG = qn("w:tab")
WORD_BREAK_TAGS = {qn("w:br"), qn("w:cr")}
WORD_RUN_TAG = qn("w:r")
WORD_RUN_PROPS_TAG = qn("w:rPr")
WORD_PARAGRAPH_PROPS_TAG = qn("w:pPr")


@dataclass
class TextSegment:
    paragraph: Paragraph
    elements: list[Any]
    text: str


def convert_docx_to_sutonny(input_data: bytes) -> tuple[BytesIO, dict[str, int]]:
    document = Document(BytesIO(input_data))
    _split_mixed_embedded_runs(document)
    _split_preserved_inline_text_runs(document)
    _enforce_english_source_fonts(document)
    _enforce_neutral_source_fonts(document)
    segments = [
        segment
        for paragraph in _iter_paragraphs(document)
        for segment in _paragraph_text_segments(paragraph)
        if segment.text and BANGLA_RE.search(segment.text)
    ]
    converted_texts = _convert_texts([segment.text for segment in segments])

    converted_segments = 0
    converted_chars = 0
    for segment, converted in _zip_equal(segments, converted_texts):
        if segment.text == converted:
            continue
        converted_chars += len(segment.text)
        _replace_text_segment(segment, converted)
        converted_segments += 1

    fallback_nodes = _remaining_bangla_text_nodes(document)
    fallback_texts = [node.text or "" for node in fallback_nodes]
    fallback_converted_texts = _convert_texts(fallback_texts)
    for node, converted in _zip_equal(fallback_nodes, fallback_converted_texts):
        if node.text == converted:
            continue
        converted_chars += len(node.text or "")
        node.text = converted
        _set_text_node_run_font(node, SUTONNY_FONT)
        converted_segments += 1

    output = BytesIO()
    document.save(output)
    output.seek(0)
    return output, {"segments": converted_segments, "characters": converted_chars}


def _convert_texts(texts: list[str]) -> list[str]:
    if not texts:
        return []
    if not FONTCONVERTER_SCRIPT.exists():
        raise RuntimeError("Bangla converter source was not found.")

    try:
        result = subprocess.run(
            ["node", str(CONVERTER_SCRIPT), str(FONTCONVERTER_SCRIPT)],
            input=json.dumps(texts, ensure_ascii=False),
            capture_output=True,
            text=True,
            encoding="utf-8",
            check=True,
            timeout=60,
        )
    except FileNotFoundError as exc:
        raise RuntimeError("Node.js is required for the bundled Bangla converter.") from exc
    except subprocess.CalledProcessError as exc:
        detail = (exc.stderr or exc.stdout or "").strip()
        raise RuntimeError(f"Bangla conversion failed. {detail}") from exc
    except subprocess.TimeoutExpired as exc:
        raise RuntimeError("Bangla conversion timed out.") from exc

    return json.loads(result.stdout)


def _zip_equal(left: list[Any], right: list[Any]) -> Iterable[tuple[Any, Any]]:
    if len(left) != len(right):
        raise RuntimeError("Bangla conversion returned an unexpected number of items.")
    return zip(left, right)


def _iter_paragraphs(document: DocumentObject) -> Iterable[Paragraph]:
    for paragraph in document.paragraphs:
        yield paragraph
    for table in document.tables:
        yield from _iter_table_paragraphs(table)

    seen_parts: set[int] = set()
    for section in document.sections:
        for part in (
            section.header,
            section.first_page_header,
            section.even_page_header,
            section.footer,
            section.first_page_footer,
            section.even_page_footer,
        ):
            marker = id(part._element)
            if marker in seen_parts:
                continue
            seen_parts.add(marker)

            for paragraph in part.paragraphs:
                yield paragraph
            for table in part.tables:
                yield from _iter_table_paragraphs(table)


def _iter_table_paragraphs(table: Table) -> Iterable[Paragraph]:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                yield paragraph
            for nested_table in cell.tables:
                yield from _iter_table_paragraphs(nested_table)


def _paragraph_text_segments(paragraph: Paragraph) -> list[TextSegment]:
    segments: list[TextSegment] = []
    current: list[Any] = []

    def flush() -> None:
        nonlocal current
        if not current:
            return
        text = "".join(_element_visible_text(element) for element in current)
        segments.append(TextSegment(paragraph=paragraph, elements=current, text=text))
        current = []

    for child in paragraph._p:
        if child.tag == WORD_PARAGRAPH_PROPS_TAG:
            continue
        if _is_convertible_text_element(child):
            current.append(child)
        else:
            flush()

    flush()
    return segments


def _is_convertible_text_element(element: Any) -> bool:
    if _contains_embedded_element(element):
        return False

    text = _element_visible_text(element)
    if not text:
        return False
    return bool(BANGLA_RE.search(text)) or _is_neutral_text(text)


def _element_visible_text(element: Any) -> str:
    parts: list[str] = []
    for node in element.iter():
        if _is_inside_embedded_element(node):
            continue
        if node.tag == WORD_TEXT_TAG:
            parts.append(node.text or "")
        elif node.tag == WORD_TAB_TAG:
            parts.append("\t")
        elif node.tag in WORD_BREAK_TAGS:
            parts.append("\n")
    return "".join(parts)


def _replace_text_segment(segment: TextSegment, text: str) -> None:
    if not segment.elements:
        return

    p_element = segment.paragraph._p
    first = segment.elements[0]
    run = _text_run_element(text)
    first.addprevious(run)

    for element in segment.elements:
        if element.getparent() is p_element:
            p_element.remove(element)


def _text_run_element(text: str) -> Any:
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    _set_rpr_font(r_pr, SUTONNY_FONT)
    run.append(r_pr)

    _append_text_to_run(run, text)
    return run


def _append_text_to_run(run: Any, text: str) -> None:
    parts = re.split(r"(\t|\n)", text)
    for part in parts:
        if not part:
            continue
        if part == "\t":
            run.append(OxmlElement("w:tab"))
            continue
        if part == "\n":
            run.append(OxmlElement("w:br"))
            continue

        text_node = OxmlElement("w:t")
        if part[:1].isspace() or part[-1:].isspace() or "  " in part:
            text_node.set(XML_SPACE, "preserve")
        text_node.text = part
        run.append(text_node)


def _set_run_font(run: Run, font_name: str) -> None:
    run.font.name = font_name
    r_pr = run._element.get_or_add_rPr()
    _set_rpr_font(r_pr, font_name)


def _split_mixed_embedded_runs(document: DocumentObject) -> None:
    for paragraph in _iter_paragraphs(document):
        for run_element in list(paragraph._p.iter(WORD_RUN_TAG)):
            if _is_inside_embedded_element(run_element):
                continue
            if not _contains_embedded_element(run_element):
                continue
            if not _element_visible_text(run_element):
                continue
            _split_run_children(run_element)


def _split_run_children(run_element: Any) -> None:
    parent = run_element.getparent()
    if parent is None:
        return

    r_pr = run_element.find(WORD_RUN_PROPS_TAG)
    new_runs: list[Any] = []
    for child in list(run_element):
        if child.tag == WORD_RUN_PROPS_TAG:
            continue

        new_run = OxmlElement("w:r")
        if r_pr is not None:
            new_run.append(deepcopy(r_pr))
        new_run.append(child)
        new_runs.append(new_run)

    for new_run in new_runs:
        run_element.addprevious(new_run)
    parent.remove(run_element)


def _split_preserved_inline_text_runs(document: DocumentObject) -> None:
    for paragraph in _iter_paragraphs(document):
        for run_element in list(paragraph._p.iter(WORD_RUN_TAG)):
            if _is_inside_embedded_element(run_element):
                continue
            if _contains_embedded_element(run_element):
                continue

            text_nodes = [child for child in run_element if child.tag == WORD_TEXT_TAG]
            other_content = [
                child
                for child in run_element
                if child.tag not in {WORD_RUN_PROPS_TAG, WORD_TEXT_TAG}
            ]
            if len(text_nodes) != 1 or other_content:
                continue

            text = text_nodes[0].text or ""
            if not BANGLA_RE.search(text) or not _has_preserved_inline_chunk(text):
                continue

            _split_text_run_by_preserved_spans(run_element, text)


def _split_text_run_by_preserved_spans(run_element: Any, text: str) -> None:
    parent = run_element.getparent()
    if parent is None:
        return

    spans = _inline_text_spans(text)
    if len(spans) < 2:
        return

    r_pr = run_element.find(WORD_RUN_PROPS_TAG)
    for span in spans:
        new_run = OxmlElement("w:r")
        if r_pr is not None:
            new_run.append(deepcopy(r_pr))
        _append_text_to_run(new_run, span)
        run_element.addprevious(new_run)
    parent.remove(run_element)


def _inline_text_spans(text: str) -> list[str]:
    spans: list[str] = []
    convert_start = 0
    for match in NON_BANGLA_CHUNK_RE.finditer(text):
        if not _should_preserve_inline_chunk(match.group(0)):
            continue
        if match.start() > convert_start:
            spans.append(text[convert_start : match.start()])
        spans.append(match.group(0))
        convert_start = match.end()
    if convert_start < len(text):
        spans.append(text[convert_start:])
    return [span for span in spans if span]


def _remaining_bangla_text_nodes(document: DocumentObject) -> list[Any]:
    nodes: list[Any] = []
    for paragraph in _iter_paragraphs(document):
        for node in paragraph._p.iter(WORD_TEXT_TAG):
            if _is_inside_embedded_element(node):
                continue
            text = node.text or ""
            if BANGLA_RE.search(text):
                nodes.append(node)
    return nodes


def _enforce_english_source_fonts(document: DocumentObject) -> None:
    for paragraph in _iter_paragraphs(document):
        for run_element in paragraph._p.iter(WORD_RUN_TAG):
            if _is_inside_embedded_element(run_element):
                continue
            if _contains_embedded_element(run_element):
                continue

            text = _element_visible_text(run_element)
            if not text or BANGLA_RE.search(text):
                continue
            if _has_english_or_math_text(text):
                _set_rpr_font(_get_or_add_rpr(run_element), ENGLISH_FONT)


def _enforce_neutral_source_fonts(document: DocumentObject) -> None:
    for paragraph in _iter_paragraphs(document):
        for run_element in paragraph._p.iter(WORD_RUN_TAG):
            if _is_inside_embedded_element(run_element):
                continue
            if _contains_embedded_element(run_element):
                continue

            text = _element_visible_text(run_element)
            if not text or BANGLA_RE.search(text):
                continue
            if _has_english_or_math_text(text):
                continue
            _set_rpr_font(_get_or_add_rpr(run_element), SUTONNY_FONT)


def _get_or_add_rpr(run_element: Any) -> Any:
    r_pr = run_element.find(WORD_RUN_PROPS_TAG)
    if r_pr is not None:
        return r_pr

    r_pr = OxmlElement("w:rPr")
    run_element.insert(0, r_pr)
    return r_pr


def _set_text_node_run_font(text_node: Any, font_name: str) -> None:
    for ancestor in text_node.iterancestors():
        if ancestor.tag == WORD_RUN_TAG:
            _set_rpr_font(_get_or_add_rpr(ancestor), font_name)
            return


def _set_rpr_font(r_pr: Any, font_name: str) -> None:
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)

    for key in ("ascii", "hAnsi", "cs", "eastAsia"):
        r_fonts.set(qn(f"w:{key}"), font_name)
    for key in ("asciiTheme", "hAnsiTheme", "csTheme", "eastAsiaTheme"):
        r_fonts.attrib.pop(qn(f"w:{key}"), None)


def _local_name(element: Any) -> str:
    tag = getattr(element, "tag", "")
    if not isinstance(tag, str):
        return ""
    return tag.rsplit("}", 1)[-1]


def _is_inside_embedded_element(element: Any) -> bool:
    for ancestor in element.iterancestors():
        local_name = _local_name(ancestor)
        if local_name in EMBEDDED_LOCAL_NAMES:
            return True
        if local_name == "p":
            return False
    return False


def _contains_embedded_element(element: Any) -> bool:
    return any(_local_name(node) in EMBEDDED_LOCAL_NAMES for node in element.iter())


def _is_word_text_token(element: Any) -> bool:
    return element.tag == WORD_TEXT_TAG or element.tag == WORD_TAB_TAG or element.tag in WORD_BREAK_TAGS


def _is_neutral_text(text: str) -> bool:
    if any(char in MATH_OPERATOR_CHARS for char in text):
        return False
    return not any(char.isalnum() for char in text)


def _has_preserved_inline_chunk(text: str) -> bool:
    return any(_should_preserve_inline_chunk(match.group(0)) for match in NON_BANGLA_CHUNK_RE.finditer(text))


def _should_preserve_inline_chunk(text: str) -> bool:
    return any(char.isascii() and char.isalnum() for char in text) or any(
        char in MATH_OPERATOR_CHARS for char in text
    )


def _has_english_or_math_text(text: str) -> bool:
    return any(char.isascii() and char.isalnum() for char in text) or any(
        char in MATH_OPERATOR_CHARS for char in text
    )
